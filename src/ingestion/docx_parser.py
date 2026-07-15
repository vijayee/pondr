"""DOCX parser -- python-docx -> ParsedDocument.

Structure-preserving: a paragraph whose style is a Word heading style (``Title``
or ``Heading 1``-``Heading 9``) opens a new section at that level; subsequent
body paragraphs append to the current section's ``content``. ``title`` = the
first ``Title``/``Heading 1`` text, else the filename stem. A doc with no
heading styles becomes one level-1 section (honest: no structure to split on).

Lazy import (GLiNER pattern): ``python-docx`` is imported inside ``parse`` so
the ingestion package stays importable without it, and a missing dep raises a
clear ``RuntimeError`` with the install hint. ``python-docx`` exposes the
top-level module name ``docx``.
"""

from __future__ import annotations

import os
import re
from typing import Optional

from .parsers import ParsedDocument, RawSection


_HEADING_RE = re.compile(r"heading\s*([1-9])", re.IGNORECASE)


def _heading_level(style_name: Optional[str]) -> Optional[int]:
    """Return the heading level (1-9) for a Word style name, else None."""
    if not style_name:
        return None
    if style_name.strip().lower() == "title":
        return 1
    m = _HEADING_RE.search(style_name)
    return int(m.group(1)) if m else None


class DocxParser:
    """``parse(source_path) -> ParsedDocument`` via python-docx."""

    source_type: str = "docx"

    def parse(self, source_path: str) -> ParsedDocument:
        try:
            import docx  # python-docx; lazy import
        except ImportError as exc:
            raise RuntimeError(
                "DocxParser needs python-docx: pip install -e .[ingestion]"
            ) from exc

        document = docx.Document(source_path)
        sections: list[RawSection] = []
        current: Optional[RawSection] = None
        title = ""

        for para in document.paragraphs:
            text = (para.text or "").strip()
            style_name = para.style.name if para.style is not None else None
            level = _heading_level(style_name)
            if level is not None and text:
                # A heading paragraph opens a new section.
                if current is not None:
                    current.content = current.content.strip()
                    if current.content or current.heading:
                        sections.append(current)
                current = RawSection(heading=text, level=level, content="")
                if not title and level == 1:
                    title = text
            else:
                if current is None:
                    # Leading body before any heading -> a root section.
                    current = RawSection(heading="", level=1, content="")
                if text:
                    current.content += text + "\n"

        if current is not None:
            current.content = current.content.strip()
            if current.content or current.heading:
                sections.append(current)

        if not title:
            title = os.path.splitext(os.path.basename(source_path))[0]

        return ParsedDocument(
            source_type=self.source_type,
            source_path=source_path,
            sections=sections,
            title=title,
        )