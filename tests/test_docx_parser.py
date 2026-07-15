"""Tests for the DOCX parser (python-docx).

The structure test is skipped cleanly when python-docx is absent
(``importorskip`` INSIDE the test, not at module level -- a module-level
``importorskip`` would skip the whole file, including the missing-dep error
test below, which is meant to run WITHOUT python-docx). The missing-dep test
needs no real .docx: ``DocxParser.parse`` lazy-imports ``docx`` and raises the
``RuntimeError`` BEFORE it opens any file.
"""

from __future__ import annotations

import pytest

from src.ingestion.docx_parser import DocxParser


def _write_docx(path: str, docx_mod) -> None:
    document = docx_mod.Document()
    document.add_heading("Project Notes", level=0)  # Title style
    document.add_paragraph("Intro about the hippocampal index.")
    document.add_heading("Alice on Storage", level=1)
    document.add_paragraph("Alice architected the storage subsystem.")
    document.add_heading("Bob on Networking", level=1)
    document.add_paragraph("Bob implemented the networking transport.")
    document.save(path)


def test_docx_headings_to_sections(tmp_path):
    pytest.importorskip("docx", reason="python-docx not installed")
    import docx as docx_mod
    fix = tmp_path / "doc.docx"
    _write_docx(str(fix), docx_mod)
    parsed = DocxParser().parse(str(fix))
    assert parsed.source_type == "docx"
    headings = [s.heading for s in parsed.sections]
    # Title (level 0 -> 1) + two Heading-1 sections.
    assert "Alice on Storage" in headings
    assert "Bob on Networking" in headings
    alice = next(s for s in parsed.sections if s.heading == "Alice on Storage")
    assert "storage" in alice.content.lower()
    # Body text is bucketed under the right section.
    bob = next(s for s in parsed.sections if s.heading == "Bob on Networking")
    assert "networking" in bob.content.lower()
    # Title taken from the Title paragraph.
    assert parsed.title == "Project Notes"


def test_docx_missing_dep_clear_error(monkeypatch):
    # No fixture needed: parse() raises at the lazy ``import docx`` BEFORE it
    # opens any file. Runs on a box WITHOUT python-docx (the box where the clear
    # error matters most).
    import builtins
    real_import = builtins.__import__

    def _no_docx(name, *a, **k):
        if name == "docx":
            raise ImportError("no docx")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", _no_docx)
    with pytest.raises(RuntimeError, match="pip install -e .\\[ingestion\\]"):
        DocxParser().parse("does-not-need-to-exist.docx")